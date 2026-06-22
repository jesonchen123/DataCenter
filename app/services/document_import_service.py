import io
import json
import re

_ALLOWED_ROLES = {"customer", "staff", "system"}
_ALLOWED_MESSAGE_FIELDS = {"role", "text", "sender"}

_ROLE_PREFIX_RE = re.compile(
    r"^(?:\[.*?\]\s*)?"                     # optional [timestamp] prefix
    r"(客户|customer|用户|user|销售|staff|客服|agent|员工|系统|system)"  # role keyword
    r"(?:[（(]([^）)]*)[）)])?"             # optional sender name in （）or ()
    r"[：:]"                                 # colon separator
)

_ROLE_MAP = {
    "客户": "customer",
    "customer": "customer",
    "用户": "customer",
    "user": "customer",
    "销售": "staff",
    "staff": "staff",
    "客服": "staff",
    "agent": "staff",
    "员工": "staff",
    "系统": "system",
    "system": "system",
}

NORMALIZE_SYSTEM_PROMPT = (
    "你是对话数据标准化助手。将输入的文本转换为客服对话消息数组。\n"
    "规则：\n"
    "1. 识别每一轮对话的说话人角色（customer=客户方提问, staff=销售/客服方回答）\n"
    "2. 忽略系统消息、时间戳、消息ID、发送人编号等元数据\n"
    "3. 每段完整发言作为一条消息，合并同一说话人的连续发言\n"
    "4. 输出必须是严格 JSON 数组，每个元素包含 role 和 text 字段，sender 字段可选：\n"
    '   [{"role":"customer"|"staff","text":"...","sender":"称呼（可选）"}]\n'
    "5. role 只能是 customer 或 staff；text 不能为空\n"
    "6. 如果文本不含有效对话，返回空数组 []\n"
    "7. 只输出 JSON 数组，不要输出其他文字或解释"
)


def extract_text(file_content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "json":
        text = file_content.decode("utf-8-sig").strip()
        return text

    if ext == "txt":
        return file_content.decode("utf-8-sig").strip()

    if ext == "docx":
        return _extract_docx_text(file_content)

    raise ValueError(f"Unsupported file format: .{ext}")


def _extract_docx_text(file_content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        )

    doc = Document(io.BytesIO(file_content))

    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                table_lines.append(" ".join(row_texts))

    all_lines = paragraphs + table_lines
    if not all_lines:
        raise ValueError("Word document contains no extractable text.")
    return "\n".join(all_lines)


def normalize_messages(raw_text: str) -> tuple[list[dict], str]:
    if not raw_text.strip():
        raise ValueError("File content is empty.")

    result = _try_json_direct(raw_text)
    if result is not None:
        return result, "json_direct"

    if _is_llm_available():
        try:
            result = _normalize_via_llm(raw_text)
            return result, "llm"
        except Exception:
            pass

    result = _normalize_via_rules(raw_text)
    return result, "rule"


def _try_json_direct(text: str) -> list[dict] | None:
    try:
        parsed = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        return None

    if isinstance(parsed, list):
        candidates = parsed
    elif isinstance(parsed, dict) and "messages" in parsed:
        candidates = parsed["messages"]
    else:
        return None

    if not isinstance(candidates, list) or not candidates:
        return None

    if not all(isinstance(item, dict) for item in candidates):
        return None

    has_role_or_text = any(
        "role" in item or "text" in item for item in candidates
    )
    if not has_role_or_text:
        return None

    return validate_message_array(candidates)


def _is_llm_available() -> bool:
    try:
        from app.services.llm_client_service import is_llm_configured

        return is_llm_configured()
    except Exception:
        return False


def _normalize_via_llm(raw_text: str) -> list[dict]:
    from app.services.llm_client_service import OpenAICompatibleLLMClient
    from app.services.llm_validation_service import extract_json_object

    client = OpenAICompatibleLLMClient()
    messages = [
        {"role": "system", "content": NORMALIZE_SYSTEM_PROMPT},
        {"role": "user", "content": raw_text},
    ]
    result = client.chat(
        messages,
        response_format={"type": "json_object"},
    )
    parsed = extract_json_object(result.content)

    if isinstance(parsed, list):
        candidates = parsed
    elif isinstance(parsed, dict) and "messages" in parsed:
        candidates = parsed["messages"]
    else:
        raise ValueError("LLM normalization output is not a message array.")

    normalized = validate_message_array(candidates)
    return normalized


def _normalize_via_rules(text: str) -> list[dict]:
    lines = text.split("\n")
    messages: list[dict] = []
    current_role: str | None = None
    current_text_parts: list[str] = []
    current_sender: str | None = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_role is not None and current_text_parts:
                _flush_message(messages, current_role, current_text_parts, current_sender)
                current_role = None
                current_text_parts = []
                current_sender = None
            continue

        match = _ROLE_PREFIX_RE.match(stripped)
        if match:
            if current_role is not None and current_text_parts:
                _flush_message(messages, current_role, current_text_parts, current_sender)

            prefix = match.group(1)
            sender_in_parens = (match.group(2) or "").strip()
            current_role = _ROLE_MAP.get(prefix, "staff")

            if sender_in_parens:
                current_sender = sender_in_parens
            else:
                current_sender = prefix if prefix != current_role else None

            content = stripped[match.end():].strip()
            current_text_parts = [content] if content else []
        else:
            if current_role is not None:
                current_text_parts.append(stripped)

    if current_role is not None and current_text_parts:
        _flush_message(messages, current_role, current_text_parts, current_sender)

    if not messages:
        raise ValueError(
            "No recognizable conversation messages found in file. "
            "Expected lines like '客户：...' or '销售：...'."
        )

    return messages


def _flush_message(messages, role, text_parts, sender):
    content = " ".join(p for p in text_parts if p).strip()
    if not content:
        return
    msg = {"role": role, "text": content}
    if sender:
        msg["sender"] = sender
    messages.append(msg)


def validate_message_array(messages: list) -> list[dict]:
    if not isinstance(messages, list) or not messages:
        raise ValueError("Messages must be a non-empty array.")

    validated: list[dict] = []
    for i, message in enumerate(messages, start=1):
        if not isinstance(message, dict):
            raise ValueError(f"Message {i} must be an object.")

        unexpected = sorted(set(message) - _ALLOWED_MESSAGE_FIELDS)
        if unexpected:
            raise ValueError(
                f"Message {i} contains unsupported fields: {', '.join(unexpected)}. "
                f"Only role, text, sender are allowed."
            )

        role = str(message.get("role") or "").strip()
        if not role:
            raise ValueError(f"Message {i}: role is required.")
        if role not in _ALLOWED_ROLES:
            raise ValueError(
                f"Message {i}: role must be customer, staff, or system, got '{role}'."
            )

        text = str(message.get("text") or "").strip()
        if not text:
            raise ValueError(f"Message {i}: text is required and must not be empty.")

        validated_msg = {"role": role, "text": text}
        sender = str(message.get("sender") or "").strip()
        if sender:
            validated_msg["sender"] = sender

        validated.append(validated_msg)

    return validated


def build_document_chat_values(messages: list[dict], metadata: dict | None = None) -> dict:
    meta = metadata or {}
    mock_chat_id = str(meta.get("mock_chat_id") or "")
    source_platform = str(meta.get("source_platform") or "document_import").strip() or "document_import"

    normalized_messages = []
    for index, msg in enumerate(messages, start=1):
        role = msg["role"]
        normalized_messages.append({
            "message_id": f"{mock_chat_id}_msg_{index:03d}",
            "sender_role": role,
            "sender_name": msg.get("sender") or _default_sender_name(role),
            "content": msg["text"],
        })

    business_line = _optional_text(meta.get("business_line"))
    product_name = _optional_text(meta.get("product_name"))
    scenario_type = _optional_text(meta.get("scenario_type"))

    raw_content = {
        "mock_chat_id": mock_chat_id,
        "source_platform": source_platform,
        "business_line": business_line,
        "product_name": product_name,
        "scenario_type": scenario_type,
        "messages": normalized_messages,
    }

    return {
        "mock_chat_id": mock_chat_id,
        "source_platform": source_platform,
        "business_line": business_line,
        "product_name": product_name,
        "scenario_type": scenario_type,
        "raw_content": raw_content,
    }


def _default_sender_name(role: str) -> str:
    if role == "customer":
        return "客户"
    if role == "staff":
        return "员工"
    return role


def _optional_text(value) -> str | None:
    text = str(value or "").strip()
    return text or None
